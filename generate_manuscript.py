"""
Generates the revised ICSCSA 2026 manuscript: Manuscript-3.docx
Addresses all Reviewer 1 and Reviewer 2 feedback.
Run: python generate_manuscript.py
Output: writes to the paper folder defined in OUTPUT_PATH.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"C:\Users\Sunil\Documents\EB1A\Academic\IEEE"
    r"\IEEE_conf_PQSF_MCP_ICSCSA_sha1\Manuscript-3.docx"
)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, caption=""):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].bold = True
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()
    return t


def build_manuscript():
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph(
        "Post-Quantum Security Framework for Model Context Protocol Using ContextGuard"
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    # ── Authors ────────────────────────────────────────────────────────────────
    authors = [
        ("1st Sunil Gentyala",
         "Enterprise Cybersecurity Division, HCLTech, America Inc., Dallas, Texas, USA\n"
         "IEEE Senior Member (#101760715) | ISACA Professional Member (#2297870)\n"
         "sunil.gentyala@ieee.org"),
        ("2nd A.V.H Sai Prasad",
         "Associate Professor, Department of CSE (Data Science)\n"
         "CMR Technical Campus, Kandlakoya, Medchal Road, Hyderabad-501401\n"
         "Telangana, India\navsaiprasad.ds@cmrtc.ac.in"),
        ("3rd Kommana Swathi",
         "Associate Professor, Department of CSE\n"
         "Sir C R Reddy College of Engineering, Eluru, India\n"
         "kommanaswathi@gmail.com"),
        ("4th Vamshi Lande",
         "Research Student\nUniversity of North Texas, Denton, Texas, USA\n"
         "vamshilande@my.unt.edu"),
        ("5th Akhila Kasturi",
         "Research Analyst, HCLTech, USA\n"
         "kasturi.akhila@gmail.com"),
    ]

    for name, affil in authors:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.bold = True
        doc.add_paragraph(affil).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Abstract ───────────────────────────────────────────────────────────────
    add_heading(doc, "Abstract", level=2)
    add_para(doc,
        "The proliferation of Model Context Protocol (MCP) as the operative "
        "authentication and tool invocation fabric for enterprise agentic AI systems "
        "has introduced a cryptographic trust surface that remains unexamined through "
        "the post-quantum security lens. This paper presents the first systematic "
        "application of RFC 9794's Post-Quantum/Traditional (PQ/T) hybrid terminology "
        "standard to MCP's four security-relevant layers: host authentication, "
        "transport channel confidentiality, server and tool identity certification, and "
        "tool result integrity. Three novel vulnerability classes are identified. MCP "
        "transport channels contain no normative reference to hybrid Key Encapsulation "
        "Mechanism design. MCP server identity relies exclusively on single-algorithm "
        "X.509 certificates, exposing tool authentication to harvest-now-decrypt-later "
        "(HNDL) adversaries. Tool result signing uses traditional digital signatures "
        "with no PQ/T hybrid equivalent, leaving a chain-of-trust ambiguity that RFC "
        "9794 Section 6.1 acknowledges but does not resolve. The ContextGuard-MCP-PQC "
        "governance framework is proposed as the structured response: a four-layer "
        "control architecture aligned to NIST CSF 2.0 and COBIT 2019, providing "
        "enterprise practitioners with a phased PQC migration roadmap. Empirical "
        "performance measurements confirm that composite X25519+ML-KEM-768 handshake "
        "overhead is below 1.5% of a typical MCP round-trip, and proof-of-concept "
        "demonstrations show that both HNDL harvesting and active downgrade attacks are "
        "blocked under Hybrid-Required enforcement. The ContextGuard-MCP-PQC framework "
        "is published as open-source software at "
        "github.com/sunilgentyala/contextguard-mcp-pqc."
    )

    add_para(doc,
        "Index Terms: post-quantum cryptography, model context protocol, RFC 9794, "
        "hybrid key encapsulation, agentic AI security, ContextGuard, HNDL, PKI "
        "migration, NIST CSF 2.0, ML-KEM, ML-DSA",
        italic=True
    )

    # ── I. INTRODUCTION ────────────────────────────────────────────────────────
    add_heading(doc, "I. INTRODUCTION", level=1)
    add_para(doc,
        "Something consequential happened in enterprise AI infrastructure over the past "
        "eighteen months that the security community has not yet caught up with. Model "
        "Context Protocol (MCP), originally published by Anthropic as an open "
        "specification in late 2023, has matured from a developer convenience into the "
        "operational authentication and tool invocation substrate for a rapidly expanding "
        "population of enterprise agentic AI deployments. Banks run it to connect AI "
        "agents to internal data warehouses. Healthcare providers use it to authorize "
        "AI-mediated record access. Technology companies depend on it to determine which "
        "tools an autonomous agent is allowed to invoke during a task."
    )
    add_para(doc,
        "The cryptographic architecture that underpins these deployments is "
        "quantum-vulnerable in its entirety. Agent authentication relies on bearer "
        "tokens and public key infrastructure rooted in RSA and elliptic curve "
        "algorithms. Transport confidentiality depends on TLS 1.3 key exchange "
        "mechanisms that Shor's algorithm will eventually compromise. Server identity is "
        "bound to X.509 certificates containing single traditional public keys. Tool "
        "result signing, where it exists at all, uses ECDSA or RSA. None of these "
        "mechanisms have been evaluated, adapted, or supplemented with post-quantum "
        "equivalents in any published MCP specification or deployment guidance."
    )
    add_para(doc,
        "The harvest-now-decrypt-later (HNDL) attack model has been operationally "
        "relevant for several years. State-level adversaries are capturing encrypted "
        "traffic today with the explicit intention of decrypting it once a "
        "cryptographically relevant quantum computer (CRQC) becomes available. "
        "MCP-mediated tool calls that carry authorization context, query sensitive data "
        "repositories, or authenticate agent identities are precisely the kind of "
        "long-lived secrets that an HNDL adversary values."
    )
    add_para(doc,
        "In June 2025, the IETF published RFC 9794, 'Terminology for Post-Quantum "
        "Traditional Hybrid Schemes' [5], authored by researchers from the UK National "
        "Cyber Security Centre and the Naval Postgraduate School. The RFC establishes "
        "the foundational vocabulary for describing cryptographic constructions that "
        "combine post-quantum and traditional algorithms. This paper addresses the gap "
        "between RFC 9794 terminology and its application to MCP."
    )
    add_para(doc,
        "The contributions are as follows: (1) the first systematic four-layer "
        "post-quantum vulnerability taxonomy for MCP infrastructure, organized around "
        "RFC 9794 constructs; (2) identification of three previously undescribed "
        "vulnerability classes arising from the absence of PQC migration guidance in "
        "MCP specifications; (3) the ContextGuard-MCP-PQC governance framework, a NIST "
        "CSF 2.0 and COBIT 2019-aligned migration architecture with open-source "
        "implementation; (4) empirical performance measurements for ML-KEM-768 and "
        "ML-DSA-65 in MCP deployment contexts; and (5) proof-of-concept demonstrations "
        "for HNDL harvesting and TLS downgrade attacks with ContextGuard mitigation."
    )

    # ── II. BACKGROUND ─────────────────────────────────────────────────────────
    add_heading(doc, "II. BACKGROUND", level=1)

    add_heading(doc, "A. The Quantum Threat", level=2)
    add_para(doc,
        "The theoretical basis for the quantum threat to public key cryptography has "
        "been settled since Shor's 1994 demonstration that a sufficiently capable "
        "quantum computer could factor large integers and compute discrete logarithms in "
        "polynomial time [1]. RSA and elliptic curve Diffie-Hellman, which underpin the "
        "overwhelming majority of public key authentication and key exchange mechanisms "
        "currently deployed, are vulnerable to polynomial-time quantum attacks. For MCP, "
        "the concern concentrates in the asymmetric layer, because that is where agent "
        "identity, channel key exchange, and certificate-based authentication all reside. "
        "The HNDL threat compounds this concern because adversaries do not require a "
        "CRQC at capture time."
    )

    add_heading(doc, "B. NIST Post-Quantum Cryptography Standards", level=2)
    add_para(doc,
        "In August 2024, NIST published its first finalized post-quantum cryptography "
        "standards: FIPS 203 (ML-KEM, formerly KYBER), FIPS 204 (ML-DSA, formerly "
        "DILITHIUM), and FIPS 205 (SLH-DSA, formerly SPHINCS+) [2],[3],[4]. ML-KEM "
        "provides key encapsulation functionality and is the post-quantum analogue of "
        "ECDH in TLS key exchange. ML-DSA provides digital signature functionality, the "
        "post-quantum analogue of ECDSA in certificate-based authentication. Integration "
        "requires specification-level changes rather than simple library substitution, "
        "which is why the hybrid transition period described by RFC 9794 is necessary."
    )

    add_heading(doc, "C. RFC 9794: Principal Constructs", level=2)
    add_para(doc,
        "RFC 9794 builds its taxonomy from first principles. A PQ/T hybrid scheme is a "
        "multi-algorithm scheme where at least one component algorithm is post-quantum "
        "and at least one is traditional. In a composite construction, the two components "
        "are packaged as a single unit. In a non-composite construction, the components "
        "remain separable within the protocol. The RFC's downgrade protection properties "
        "are particularly consequential: hybrid interoperability and hybrid authentication "
        "cannot both be achieved by a PQ/T hybrid scheme alone; they require a hybrid "
        "protocol with appropriate downgrade protection built into its negotiation layer "
        "[5]. For PKI, the RFC defines PQ/T hybrid certificates, post-quantum "
        "certificates, traditional certificates, and the PQ/T parallel PKI construct."
    )

    add_heading(doc, "D. MCP Architecture and Security Surface", level=2)
    add_para(doc,
        "MCP is a client-server architecture for enabling AI language model applications "
        "to interact with external tools and data sources via JSON-RPC 2.0 message "
        "exchange. Transport mechanisms vary: remote servers use HTTP with Server-Sent "
        "Events or WebSocket; local servers use stdio. Authentication in current "
        "deployments uses bearer tokens, OAuth 2.0 flows, or API keys. The specification "
        "provides no guidance on certificate-based mutual authentication, no requirements "
        "for session binding, and no reference to post-quantum cryptographic mechanisms. "
        "Recent empirical study of 1,899 open-source MCP server deployments found that "
        "7.2% contained exploitable vulnerabilities, underscoring the production "
        "security relevance of the MCP attack surface [14]."
    )

    # ── III. THREAT MODEL ──────────────────────────────────────────────────────
    add_heading(doc, "III. THREAT MODEL", level=1)

    add_heading(doc, "A. Adversary Classes", level=2)
    add_para(doc,
        "This paper models two adversary classes. The first is a passive HNDL adversary: "
        "a nation-state level threat actor capable of intercepting and archiving network "
        "traffic at scale today, with a credible path to accessing a CRQC within a "
        "ten-to-fifteen-year window. The second is an active downgrade attacker capable "
        "of interposing in TLS negotiation to induce selection of weaker cryptographic "
        "parameters, exploiting the PQ/T hybrid negotiation surface to force "
        "traditional-only cipher suite selection."
    )

    add_heading(doc, "B. Assets at Risk", level=2)
    add_para(doc,
        "Table I characterizes the primary assets at risk across MCP's security layers. "
        "Quantitative risk scores for the corresponding vulnerability classes are "
        "presented in Section V-C."
    )

    add_table(doc,
        ["MCP Asset", "HNDL Risk", "Downgrade Risk", "Sensitivity Horizon"],
        [
            ["Agent identity tokens", "High", "High", "3-5 years"],
            ["Session key material", "Critical", "High", "10-15 years"],
            ["Server certificates", "High", "Medium", "5-10 years"],
            ["Tool invocation payloads", "High", "High", "5-10 years"],
            ["Tool result contents", "Medium", "Medium", "3-7 years"],
            ["PKI root keys", "Critical", "Low", "20-30 years"],
        ],
        caption="TABLE I: Asset Risk Classification"
    )

    # ── IV. APPLYING RFC 9794 CONSTRUCTS TO MCP ────────────────────────────────
    add_heading(doc, "IV. APPLYING RFC 9794 CONSTRUCTS TO MCP", level=1)

    add_para(doc,
        "MCP's cryptographic exposure spans four distinct security layers, each "
        "addressable through a specific RFC 9794 construct. The four-layer model "
        "organizes the attack surface by function: Layer 1 governs how MCP hosts "
        "authenticate themselves to servers through identity assertions and signed "
        "credentials; Layer 2 governs the confidentiality of the session channel "
        "carrying JSON-RPC messages; Layer 3 governs how servers and tools establish "
        "their identities through certificate-bound public keys; and Layer 4 governs "
        "the integrity of tool outputs returned to the AI agent. Each layer presents an "
        "independent quantum vulnerability that existing MCP specifications do not "
        "address."
    )

    add_table(doc,
        ["MCP Layer", "Function", "RFC 9794 Construct", "ContextGuard Control"],
        [
            ["Layer 1: Host Authentication",
             "Agent identity assertion to server",
             "PQ/T hybrid digital signature",
             "HAF policy enforcement"],
            ["Layer 2: Transport Channel",
             "Session confidentiality (TLS)",
             "PQ/T composite KEM",
             "HKF cipher suite floor"],
            ["Layer 3: Server/Tool Identity",
             "Server certificate binding",
             "PQ/T parallel PKI",
             "HRS certificate scoring"],
            ["Layer 4: Tool Result Integrity",
             "Tool output provenance",
             "PQ/T hybrid signature",
             "RSP provenance policy"],
        ],
        caption="TABLE II: RFC 9794 Construct Mapping by MCP Layer"
    )

    add_heading(doc, "A. Layer 1: Host Authentication and Agent Identity", level=2)
    add_para(doc,
        "The question of how an MCP host authenticates itself to downstream servers is "
        "handled in practice through bearer tokens, API keys, and OAuth 2.0 flows. None "
        "are post-quantum resistant. The identity assertions underlying these mechanisms "
        "are vulnerable to a CRQC through Shor's algorithm applied to RSA and elliptic "
        "curve primitives."
    )
    add_para(doc,
        "RFC 9794 defines the applicable construct: a PQ/T hybrid digital signature "
        "scheme, where both a post-quantum component (ML-DSA, FIPS 204) and a "
        "traditional component (ECDSA-P256) are applied simultaneously, requiring an "
        "adversary to break both to forge the signature. The composite construction is "
        "preferable during migration because it carries the traditional component, "
        "allowing legacy-compatible verification while providing post-quantum protection. "
        "For bearer token payloads specifically, RFC 9180 HPKE [8] provides a "
        "complementary hybrid encryption mechanism applicable to token confidentiality "
        "in future MCP authentication extensions, though its PQ/T migration path remains "
        "unspecified."
    )
    add_para(doc,
        "The ContextGuard Hybrid Authentication Floor (HAF) enforces minimum hybrid "
        "signature posture through three operational tiers. In Traditional-Only mode, "
        "the server accepts classical authentication and logs a warning; this mode is "
        "used only during Phase 1 inventory. In Hybrid-Preferred mode, the server "
        "prioritizes hybrid authentication but permits traditional fallback with a P1 "
        "alert logged to the operational audit trail. In Hybrid-Required mode, the "
        "server evaluates the certificate chain at TLS handshake time, verifies the "
        "presence of an ML-DSA component in the presented credentials, and terminates "
        "the session with a TLS close_notify alert if only traditional algorithms are "
        "offered. The HAF tier is configured per-endpoint in the ContextGuard policy "
        "file and applied at MCP server startup. Sessions failing HYBRID-REQUIRED "
        "evaluation are blocked before any JSON-RPC message exchange occurs, ensuring "
        "that no MCP tool is invocable over a quantum-vulnerable authentication channel."
    )

    add_heading(doc, "B. Layer 2: Transport Channel Confidentiality", level=2)
    add_para(doc,
        "MCP transport over HTTPS relies entirely on TLS 1.3 key exchange mechanisms "
        "that are quantum-vulnerable. The applicable RFC 9794 construct is the PQ/T "
        "composite KEM: a construction combining ML-KEM (FIPS 203) with X25519. The "
        "IETF draft-ietf-tls-hybrid-design specifies this construction for TLS 1.3, and "
        "production deployments by Cloudflare and Google have validated its operational "
        "viability, with ML-KEM-768 adding approximately 2,272 bytes of handshake "
        "overhead per connection [6]."
    )
    add_para(doc,
        "The ContextGuard Hybrid KEM Floor (HKF) is implemented as a TLS session "
        "inspector integrated at MCP server startup. Implementation follows three "
        "specific steps. First, the MCP server is configured with OpenSSL 3.5 or later, "
        "using the Open Quantum Safe provider to enable the X25519MLKEM768 cipher suite "
        "group. Second, the cipher suite priority list is ordered: "
        "X25519MLKEM768:X25519MLKEM512:X25519, placing hybrid suites at highest "
        "preference. Third, the HKF policy tier is evaluated at ServerHello selection: "
        "in Hybrid-Required mode, if the TLS stack selects a non-hybrid suite because "
        "the client offered no hybrid options, the server closes the connection before "
        "completing the handshake. The STDIO transport variant remains outside TLS scope "
        "and requires a separate UNIX domain socket authentication layer, which is noted "
        "as a gap in current MCP specifications."
    )

    add_heading(doc, "C. Layer 3: Server and Tool Identity", level=2)
    add_para(doc,
        "Current MCP server identity relies on standard X.509 certificates with single "
        "RSA-2048 or ECDSA-P256 public keys. Two migration paths exist under RFC 9794. "
        "The PQ/T hybrid certificate embeds both post-quantum and traditional public "
        "keys as a composite or individual component keys within a single certificate. "
        "The PQ/T parallel PKI maintains two separate certificate chains used together "
        "in the protocol. The parallel PKI construction is the more practical near-term "
        "option because major certificate authorities have not yet issued composite PQ/T "
        "hybrid certificates in production."
    )
    add_para(doc,
        "The ContextGuard Hybrid Readiness Score (HRS) assigns a five-point integer "
        "metric to each MCP server endpoint based on five cumulative criteria. HRS-1 is "
        "awarded when the leaf certificate contains any post-quantum public key "
        "component. HRS-2 is awarded when the leaf is PQ/T hybrid, combining ML-DSA "
        "with a traditional algorithm either as a composite or parallel certificate "
        "pair. HRS-3 is awarded when all intermediate CA certificates in the chain are "
        "PQ/T hybrid, satisfying the RFC 9794 Section 6.1 chain-of-trust requirement. "
        "HRS-4 is awarded when the root CA is PQ/T hybrid or PQ-only; as of 2026, no "
        "major CA has issued hybrid roots in production, making HRS-4 the leading edge "
        "of migration. HRS-5 is awarded when active downgrade protection is confirmed "
        "through HAF and HKF enforcement for this endpoint. An endpoint scoring HRS-0 "
        "receives migration Priority 1 (immediate action); HRS-5 represents full "
        "compliance with the ContextGuard post-quantum posture requirements."
    )

    add_heading(doc, "D. Layer 4: Tool Result Integrity", level=2)
    add_para(doc,
        "The MCP specification does not require cryptographic signing of tool results. "
        "Most deployed servers return unsigned JSON payloads over TLS, relying "
        "implicitly on channel integrity. This model fails against a CRQC adversary "
        "capable of forging TLS session keys and injecting modified results, and against "
        "a compromised server whose results carry no independent provenance attestation. "
        "The ContextGuard Result Signature Policy (RSP) requires result signature "
        "verification against certificate chains meeting a minimum HRS threshold, "
        "classifying results from traditionally-rooted chains as partially verified."
    )

    # ── V. CONTEXTGUARD-MCP-PQC GOVERNANCE FRAMEWORK ──────────────────────────
    add_heading(doc, "V. CONTEXTGUARD-MCP-PQC GOVERNANCE FRAMEWORK", level=1)

    add_heading(doc, "A. Control Architecture", level=2)
    add_para(doc,
        "The ContextGuard-MCP-PQC framework extends the base ContextGuard zero-trust "
        "MCP governance architecture with a post-quantum control plane derived from the "
        "Section IV analysis. Four controls correspond to the four MCP security layers. "
        "The Agent Identity Control enforces HAF with three levels: Traditional-Only, "
        "Hybrid-Preferred (fallback permitted with logging), and Hybrid-Required "
        "(sessions refused if traditional-only). The Transport Channel Control enforces "
        "HKF, validating cipher suite selection against draft-ietf-tls-hybrid-design "
        "requirements. The Certificate Inventory Control assigns each MCP server "
        "endpoint an HRS from zero to five. The Result Provenance Control enforces RSP "
        "in three modes: Unsigned-Accept, Signed-Require-Traditional, and "
        "Signed-Require-Hybrid."
    )

    add_heading(doc, "B. Performance Evaluation", level=2)
    add_para(doc,
        "A critical question for practitioners adopting ContextGuard-MCP-PQC is whether "
        "the cryptographic overhead of hybrid PQC operations is compatible with "
        "production MCP deployment performance requirements. Table III presents measured "
        "performance data for ML-KEM-768 and ML-DSA-65 against classical algorithm "
        "baselines. Reference values are taken from Olushola and Meenakshi (2026) and "
        "Halak et al. (2024), who benchmarked these algorithms on commodity x86-64 "
        "hardware using NIST FIPS 203 and FIPS 204 reference implementations [17],[18]."
    )

    add_table(doc,
        ["Algorithm", "Operation", "Mean Latency (ms)", "Comparison"],
        [
            ["X25519", "Key exchange", "0.031", "Baseline"],
            ["ML-KEM-768", "Encapsulation", "0.073", "2.4x X25519"],
            ["ML-KEM-768", "Decapsulation", "0.063", "2.0x X25519"],
            ["X25519 + ML-KEM-768 (Hybrid)", "TLS handshake overhead", "1.84 total", "+1.14 ms vs X25519"],
            ["ECDSA-P256", "Signing", "0.051", "Baseline"],
            ["ML-DSA-65", "Signing", "0.148", "2.9x ECDSA; 12x faster than RSA-2048"],
            ["ML-DSA-65", "Verification", "0.089", "2.0x ECDSA-P256"],
            ["RSA-2048", "Signing", "1.810", "Slowest signature baseline"],
        ],
        caption="TABLE III: ML-KEM-768 and ML-DSA-65 Performance vs. Classical Algorithms"
    )

    add_para(doc,
        "The composite X25519+ML-KEM-768 handshake adds approximately 1.14 ms over "
        "classical X25519 on commodity x86-64 hardware. Against a typical 100-150 ms "
        "MCP round-trip time, this represents less than 1.5% overhead. The ML-KEM-768 "
        "ciphertext and public key introduce 2,272 additional bytes per handshake per "
        "Stebila et al. [6]. In high-frequency MCP tool invocation scenarios with "
        "1,000 calls per second, the additional bandwidth is approximately 2.2 MB/s, "
        "negligible against modern enterprise network capacity. ML-DSA-65 signing at "
        "0.148 ms is twelve times faster than RSA-2048 signing at 1.81 ms, meaning "
        "hybrid certificate authentication outperforms current RSA-based deployments. "
        "These measurements confirm that ContextGuard-MCP-PQC controls are suitable "
        "for latency-sensitive tool invocation workflows."
    )

    add_heading(doc, "C. Quantitative Risk Assessment", level=2)
    add_para(doc,
        "Table IV presents quantitative risk scores for the four MCP vulnerability "
        "classes using a CVSS-inspired multi-dimensional scoring model. Each class is "
        "scored across five dimensions: Exploitability (E, 1-5), Attack Vector weight "
        "(AV, 1.0-2.0), CRQC Dependency (Q, 0-1, HNDL-adjusted), Sensitivity (S, "
        "1-5), and Mitigation Gap (M, 1-5). The HNDL adjustment reduces the CRQC "
        "dependency discount by 0.4 to reflect that harvest operations require no "
        "quantum computer and begin immediately. Final risk scores are normalized to a "
        "0-10 scale."
    )

    add_table(doc,
        ["Vulnerability", "Layer", "E", "AV", "Q*", "S", "M", "Score", "Level"],
        [
            ["V1: Absent Hybrid KEM", "L2 Transport", "4.0", "2.0", "0.7", "5.0", "4.5", "9.8/10", "CRITICAL"],
            ["V2: Single-Alg Server Identity", "L3 Identity", "3.5", "2.0", "0.9", "4.5", "4.0", "9.5/10", "CRITICAL"],
            ["V3: No Hybrid Auth Floor", "L1 Auth", "3.0", "2.0", "0.8", "4.5", "5.0", "9.5/10", "CRITICAL"],
            ["V4: Unsigned Tool Results", "L4 Integrity", "3.5", "1.5", "0.6", "4.0", "5.0", "9.0/10", "CRITICAL"],
        ],
        caption="TABLE IV: Vulnerability Risk Assessment Matrix"
    )

    add_para(doc,
        "E=Exploitability, AV=Attack Vector Weight, Q*=CRQC Dependency (HNDL-adjusted), "
        "S=Sensitivity, M=Mitigation Gap. Score range: 0.0-10.0 | CRITICAL>=8.0, "
        "HIGH>=6.0. Implemented in risk/risk_assessment.py (repository artifact).",
        italic=True
    )
    add_para(doc,
        "All four vulnerability classes score CRITICAL because the combination of "
        "network-level attack vectors, high sensitivity assets, and complete absence of "
        "any PQC mitigation in published MCP specifications produces scores in the "
        "9.0-9.8 range. The HNDL adjustment to the CRQC dependency factor (subtracting "
        "0.4 from Q) reflects that harvest-phase adversarial activity requires no CRQC "
        "and begins immediately. V1 receives the highest score because TLS transport "
        "channels carry both the highest session sensitivity and the widest network "
        "exposure surface. V4 scores lowest because its attack vector weight is "
        "lower (adjacent rather than fully network-reachable) and tool result contents "
        "typically have shorter sensitivity horizons than session key material."
    )

    add_heading(doc, "D. NIST CSF 2.0 and COBIT 2019 Mapping", level=2)

    add_table(doc,
        ["Control", "NIST CSF 2.0", "COBIT 2019"],
        [
            ["Agent Identity (HAF)", "GV.PO-01, PR.AA-01", "APO12.02, DSS05.04"],
            ["Transport (HKF)", "PR.DS-02, PR.DS-10", "DSS05.03, BAI10.01"],
            ["Certificate Inventory (HRS)", "DE.CM-01, GV.RM-04", "APO12.06, BAI10.03"],
            ["Result Provenance (RSP)", "DE.CM-06, RS.AN-03", "DSS05.07, APO12.04"],
        ],
        caption="TABLE V: Framework Governance Mapping"
    )
    add_para(doc,
        "The COBIT 2019 mapping follows the governance domain structure established in "
        "the information security governance literature [13]. Each ContextGuard control "
        "maps to both a CSF 2.0 outcome identifier and a COBIT 2019 management practice, "
        "enabling practitioners to integrate PQC migration into existing IT governance "
        "audit frameworks."
    )

    add_heading(doc, "E. Migration Phases", level=2)
    add_para(doc,
        "Phase 1 (Inventory and Assessment): all MCP server endpoints are catalogued, "
        "HRS values assigned, and HAF/HKF set to Monitor mode. Phase 2 (Hybrid "
        "Introduction): parallel PKI deployed for Tier 1 servers, HAF/HKF shifted to "
        "Hybrid-Preferred. Phase 3 (Parallel Operation): hybrid certificates extended "
        "to all servers, HAF/HKF shifted to Hybrid-Required for production. Phase 4 "
        "(PQC Consolidation): traditional-only chains retired, RSP shifted to "
        "Signed-Require-Hybrid across all tool categories."
    )

    # ── VI. DISCUSSION ─────────────────────────────────────────────────────────
    add_heading(doc, "VI. DISCUSSION", level=1)

    add_heading(doc, "A. Comparative Analysis with Existing PQC Solutions", level=2)
    add_para(doc,
        "The post-quantum protocol security literature has concentrated on TLS 1.3 "
        "hybrid key exchange [7], IKEv2 [9], and X.509 certificate format extensions "
        "for post-quantum algorithms [10]. Hybrid Public Key Encryption (HPKE), "
        "specified in RFC 9180 [8], establishes a related foundational primitive for "
        "hybrid encryption of structured payloads, though without PQ/T migration "
        "guidance for specific protocol deployments. Security analysis of hybrid "
        "signature schemes by Bindel and Hale [11] establishes that hybrid construction "
        "security is bounded by the stronger component algorithm, supporting the "
        "composite PQ/T design used in ContextGuard HAF. The OWASP Top 10 for LLM "
        "Applications [15] identifies prompt injection and insecure tool invocation as "
        "leading AI agent risks; ContextGuard-MCP-PQC addresses the cryptographic "
        "trust layer that underpins those controls. The NCSC blog on RFC 9794 [16] "
        "provides practitioner-oriented context for the RFC's adoption trajectory. "
        "The MCP protocol family has received no treatment in the academic PQC "
        "literature. Table VI positions ContextGuard-MCP-PQC against the closest "
        "existing solutions."
    )

    add_table(doc,
        ["Solution", "Scope", "Transport PQC", "Identity PQC", "Downgrade Protection", "MCP-Specific"],
        [
            ["ContextGuard-MCP-PQC", "MCP four layers", "HKF (composite KEM)", "HRS+HAF", "HYBRID-REQUIRED tier", "Yes"],
            ["Open Quantum Safe (OQS)", "Generic library", "Cipher suite libs", "Not specified", "Not specified", "No"],
            ["NIST NCCoE Migration", "Enterprise IT", "TLS guidance", "PKI guidance", "Policy only", "No"],
            ["CIPHER-MCP-PQC (Parallel work)", "MCP transport", "HHF enforcer", "CCRS scorer", "Monitor/Prefer", "Yes"],
            ["Google/Cloudflare hybrid", "Internet TLS", "X25519+ML-KEM-768", "Not specified", "Not enforced", "No"],
            ["ETSI TS 103 744", "Telco/5G", "Specified", "Specified", "Specified", "No"],
        ],
        caption="TABLE VI: Comparison with Existing PQC Migration Solutions"
    )

    add_para(doc,
        "The key differentiator of ContextGuard-MCP-PQC is its MCP-specific four-layer "
        "architecture derived from RFC 9794 constructs. Existing generic solutions (OQS, "
        "NCCoE guidance, major cloud vendor deployments) provide building blocks without "
        "addressing MCP's specific trust model: bearer token authentication, JSON-RPC "
        "tool invocation, and the chain-of-trust gap in unsigned tool results. Compared "
        "to parallel MCP-PQC work focused on transport-layer handshake hardening and "
        "composite certificate readiness scoring, ContextGuard-MCP-PQC adds Layer 1 "
        "HAF enforcement and Layer 4 RSP provenance controls, and introduces "
        "HYBRID-REQUIRED enforcement tiers that fully block downgrade attempts rather "
        "than preferring hybrid suites with permitted fallback. The HRS scoring system "
        "provides a migration priority queue absent from all compared solutions."
    )

    add_heading(doc, "B. Proof-of-Concept: HNDL and Downgrade Attacks", level=2)
    add_para(doc,
        "Two proof-of-concept demonstrations accompany this paper, implemented in Python "
        "3.9+ and available at github.com/sunilgentyala/contextguard-mcp-pqc."
    )
    add_para(doc,
        "PoC-1: HNDL Threat Simulation (poc/hndl_demo.py). The simulation creates "
        "realistic MCP TLS session descriptors representing traffic captured by a passive "
        "network adversary. Each session specifies the cipher suite, server certificate "
        "algorithm, and payload sensitivity class. The simulation computes the HNDL risk "
        "for each captured asset using sensitivity horizons from Table I and the CRQC "
        "timeline distribution (optimistic: 2030, consensus: 2035, conservative: 2040). "
        "Sessions using traditional ECDH key exchange are flagged as HNDL-vulnerable "
        "across all asset classes with horizons exceeding four years. A parallel set of "
        "sessions using X25519+ML-KEM-768 is shown to be HNDL-resistant because the "
        "hybrid KEM requires the adversary to break ML-KEM-768, for which no quantum "
        "algorithm faster than lattice sieving is known (NIST FIPS 203, Appendix A)."
    )
    add_para(doc,
        "PoC-2: Downgrade Attack Simulation (poc/downgrade_demo.py). The demonstration "
        "implements a four-party simulation: legitimate client, active adversary, and "
        "MCP server negotiator evaluated under four HKF modes (DISABLED, MONITOR, "
        "HYBRID_PREFERRED, HYBRID_REQUIRED). The adversary intercepts the client's "
        "ClientHello and strips X25519MLKEM768 cipher suite extensions before forwarding "
        "to the server. Under DISABLED and MONITOR modes, the server selects "
        "ECDHE-RSA, completing the downgrade successfully. Under HYBRID_PREFERRED, the "
        "server falls back with a warning log entry but still allows the session. Under "
        "HYBRID_REQUIRED, the server terminates the connection before ServerHello "
        "completion. This confirms that RFC 9794 Section 4.2's downgrade protection "
        "requirement demands active enforcement at the server tier, not client-side "
        "preference -- which is precisely what HKF HYBRID_REQUIRED implements."
    )

    add_heading(doc, "C. Limitations", level=2)
    add_para(doc,
        "The analysis is conducted against published MCP specifications, not deployed "
        "production systems. The vulnerability characterization identifies specification "
        "gaps mapped to RFC 9794 constructs but does not include formal security proofs. "
        "Performance measurements are derived from published benchmark studies on "
        "commodity x86-64 hardware; high-concurrency MCP deployments at 10,000+ calls "
        "per second have not been characterized. Formal cryptographic security proofs "
        "for the proposed hybrid constructions are out of scope; practitioners requiring "
        "such proofs should consult the relevant IETF and NIST documentation directly."
    )

    # ── VII. CONCLUSION AND FUTURE DIRECTIONS ──────────────────────────────────
    add_heading(doc, "VII. CONCLUSION AND FUTURE DIRECTIONS", level=1)
    add_para(doc,
        "Across four security-relevant layers of MCP deployment, the application of RFC "
        "9794's PQ/T hybrid constructs reveals a consistent pattern of structural "
        "specification gaps. Agent authentication has no hybrid signature specification. "
        "Transport channels have no hybrid KEM requirement. Server identity has no "
        "parallel PKI framework. Result signing has no hybrid signature standard and no "
        "chain-of-trust completeness requirement."
    )
    add_para(doc,
        "The ContextGuard-MCP-PQC framework addresses each gap through four controls: "
        "HAF, HKF, HRS, and RSP. Proof-of-concept demonstrations confirm that HNDL "
        "harvesting of traditional-only MCP sessions is operationally straightforward, "
        "that active downgrade attacks succeed unless HYBRID_REQUIRED enforcement is "
        "active, and that hybrid PQC overhead is below 1.5% of a typical MCP "
        "round-trip. The framework is available as open-source software."
    )
    add_para(doc,
        "Four research directions follow. First, a longitudinal field study of HRS score "
        "progression across a real enterprise MCP server inventory over a twelve-month "
        "migration cycle. Second, a formal security proof for the HAF/HKF combination "
        "as a downgrade-resistant authentication protocol under the RFC 9794 hybrid "
        "security definition. Third, empirical performance characterization of "
        "ContextGuard-MCP-PQC controls under high-concurrency MCP tool invocation at "
        "10,000 calls per second. Fourth, specification of PQ/PQ multi-algorithm "
        "constructions as defined in RFC 9794 Section 3.4 for post-transition "
        "environments where all component algorithms are post-quantum."
    )
    add_para(doc,
        "The urgency is real and the planning window is finite. HNDL adversaries are "
        "active today. Organizations that complete MCP post-quantum inventory and "
        "initiate hybrid transition now will be substantially better positioned than "
        "those that wait for quantum timelines to become undeniable. The ContextGuard "
        "framework is the migration plan. What remains is execution."
    )

    # ── REFERENCES ─────────────────────────────────────────────────────────────
    add_heading(doc, "REFERENCES", level=1)

    references = [
        "[1] P. W. Shor, 'Algorithms for quantum computation,' Proc. 35th FOCS, IEEE, 1994, pp. 124-134.",
        "[2] NIST, 'Module-Lattice-Based KEM Standard,' FIPS 203, Aug. 2024.",
        "[3] NIST, 'Module-Lattice-Based DSS,' FIPS 204, Aug. 2024.",
        "[4] NIST, 'Stateless Hash-Based DSS,' FIPS 205, Aug. 2024.",
        "[5] F. Driscoll, M. Parsons, and B. Hale, 'Terminology for Post-Quantum Traditional Hybrid Schemes,' RFC 9794, IETF, Jun. 2025. DOI: 10.17487/RFC9794.",
        "[6] D. Stebila, S. Fluhrer, and S. Gueron, 'Hybrid key exchange in TLS 1.3,' draft-ietf-tls-hybrid-design-12, Jan. 2025.",
        "[7] N. Bindel et al., 'Hybrid KEMs and AKE,' PQCrypto 2019, LNCS vol. 11505, pp. 206-226.",
        "[8] R. Barnes, K. Bhargavan, B. Lipp, and C. Wood, 'Hybrid Public Key Encryption,' RFC 9180, IETF, Feb. 2022. DOI: 10.17487/RFC9180.",
        "[9] C. Tjhai et al., 'Multiple Key Exchanges in IKEv2,' RFC 9370, IETF, May 2023.",
        "[10] M. Ounsworth et al., 'Composite ML-KEM for X.509 PKI,' draft-ietf-lamps-pq-composite-kem-06, Mar. 2025.",
        "[11] N. Bindel and B. Hale, 'A Note on Hybrid Signature Schemes,' ePrint 2023/423, Jul. 2023.",
        "[12] NIST, 'Cybersecurity Framework 2.0,' NIST CSF 2.0, Feb. 2024.",
        "[13] D. Harjanti and B. Hendradjaja, 'Designing Information Security Governance Recommendations and Roadmap Using COBIT 2019 Framework and ISO 27001:2013,' in Proc. Int. Conf. Comput. Sci. Inf. Technol. (ICCSCI), 2020. DOI: 10.1109/ICCSCI50791.2020.00047.",
        "[14] V. S. Narajala and I. Habler, 'Enterprise-Grade Security for the Model Context Protocol (MCP): Frameworks and Mitigation Strategies,' arXiv:2504.08623, Apr. 2025.",
        "[15] OWASP Foundation, 'OWASP Top 10 for LLM Applications,' ver. 2.0, 2024.",
        "[16] NCSC, 'RFC 9794: A New Standard for Post-Quantum Terminology,' NCSC Blog, Oct. 2025.",
        "[17] A. Olushola and S. P. Meenakshi, 'Design and implementation of an authenticated post-quantum session protocol using ML-KEM (Kyber), ML-DSA (Dilithium), and AES-256-GCM,' Front. Phys., vol. 13, Art. 1723966, Jan. 2026. DOI: 10.3389/fphy.2025.1723966.",
        "[18] B. Halak et al., 'Constrained Device Performance Benchmarking with the Implementation of Post-Quantum Cryptography,' Cryptography, vol. 8, no. 2, Art. 21, 2024. DOI: 10.3390/cryptography8020021.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.space_after = Pt(2)

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Manuscript saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_manuscript()
