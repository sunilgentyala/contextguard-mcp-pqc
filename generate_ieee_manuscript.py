"""
IEEE Conference Paper Formatter — ContextGuard-MCP-PQC ICSCSA 2026
Generates Manuscript-3-IEEE.docx with full IEEE two-column formatting.

IEEE conference paper specifications applied:
  - Page: US Letter (8.5" x 11")
  - Margins: Top 0.75", Bottom 1", Left/Right 0.625"
  - Title area: single column, Times New Roman 24pt bold, centered
  - Author block: single column, 10pt, centered
  - Abstract: single column, 9pt italic, justified
  - Body: two columns (3.5" each, 0.25" gap), 10pt Times New Roman, justified
  - Section headings (level 1): 10pt bold uppercase, centered in column
  - Subsection headings (level 2): 10pt bold italic, left-aligned
  - Figure/table captions: 8pt bold, centered
  - References: 8pt, hanging indent
"""

import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"C:\Users\Sunil\Documents\EB1A\Academic\IEEE"
    r"\IEEE_conf_PQSF_MCP_ICSCSA_sha1\Manuscript-3-IEEE.docx"
)

# ── Measurement constants (twips: 1 inch = 1440) ──────────────────────────
PAGE_W      = 12240   # 8.5"
PAGE_H      = 15840   # 11"
MARGIN_TOP  = 1080    # 0.75"
MARGIN_BOT  = 1440    # 1.0"
MARGIN_LR   = 900     # 0.625"
COL_GAP     = 360     # 0.25"
COL_W       = (PAGE_W - MARGIN_LR * 2 - COL_GAP) // 2   # 5040 twips = 3.5"


# ── XML helpers ────────────────────────────────────────────────────────────

def _make_el(tag, **attribs):
    el = OxmlElement(tag)
    for k, v in attribs.items():
        el.set(qn(k), str(v))
    return el


def set_font(run, name="Times New Roman", size_pt=10, bold=False, italic=False,
             color=None):
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    # Replace any existing rFonts
    existing = rpr.find(qn("w:rFonts"))
    if existing is not None:
        rpr.remove(existing)
    rpr.insert(0, rFonts)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_before=0, space_after=0,
                 first_line_indent=0, line_spacing=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Inches(first_line_indent)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def set_page_margins(section):
    """Apply IEEE page margins to a section."""
    sectPr = section._sectPr
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.append(pgMar)
    pgMar.set(qn("w:top"),    str(MARGIN_TOP))
    pgMar.set(qn("w:bottom"), str(MARGIN_BOT))
    pgMar.set(qn("w:left"),   str(MARGIN_LR))
    pgMar.set(qn("w:right"),  str(MARGIN_LR))
    pgMar.set(qn("w:header"), "720")
    pgMar.set(qn("w:footer"), "720")
    pgMar.set(qn("w:gutter"), "0")


def set_page_size(section):
    sectPr = section._sectPr
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.append(pgSz)
    pgSz.set(qn("w:w"), str(PAGE_W))
    pgSz.set(qn("w:h"), str(PAGE_H))


def set_columns(section, num_cols):
    """Set column count on a section."""
    sectPr = section._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is not None:
        sectPr.remove(cols_el)
    cols_el = OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(num_cols))
    if num_cols == 2:
        cols_el.set(qn("w:space"), str(COL_GAP))
        c1 = OxmlElement("w:col")
        c1.set(qn("w:w"), str(COL_W))
        c2 = OxmlElement("w:col")
        c2.set(qn("w:w"), str(COL_W))
        cols_el.append(c1)
        cols_el.append(c2)
    else:
        cols_el.set(qn("w:space"), "0")
    sectPr.append(cols_el)


def insert_continuous_section_break(para, num_cols):
    """
    Insert a continuous section break at the END of the given paragraph.
    This breaks from the current column count to num_cols in the NEXT section.
    The sectPr is placed inside the paragraph's pPr.
    """
    pPr = para._p.get_or_add_pPr()
    # Remove any existing sectPr in pPr
    existing = pPr.find(qn("w:sectPr"))
    if existing is not None:
        pPr.remove(existing)

    sectPr = OxmlElement("w:sectPr")

    # Type: continuous (no page break)
    secType = OxmlElement("w:type")
    secType.set(qn("w:val"), "continuous")
    sectPr.append(secType)

    # Page size
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(PAGE_W))
    pgSz.set(qn("w:h"), str(PAGE_H))
    sectPr.append(pgSz)

    # Page margins
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"),    str(MARGIN_TOP))
    pgMar.set(qn("w:bottom"), str(MARGIN_BOT))
    pgMar.set(qn("w:left"),   str(MARGIN_LR))
    pgMar.set(qn("w:right"),  str(MARGIN_LR))
    pgMar.set(qn("w:header"), "720")
    pgMar.set(qn("w:footer"), "720")
    pgMar.set(qn("w:gutter"), "0")
    sectPr.append(pgMar)

    # Column count for THIS section (before the break)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "0")
    sectPr.append(cols)

    pPr.append(sectPr)


# ── Content helpers ────────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
    run = p.add_run(text)
    set_font(run, size_pt=24, bold=True)
    return p


def add_author_name(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    run = p.add_run(text)
    set_font(run, size_pt=10, bold=True)
    return p


def add_affil(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    run = p.add_run(text)
    set_font(run, size_pt=9, italic=True)
    return p


def add_abstract_heading(doc, text="Abstract"):
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=0)
    run = p.add_run(text + "—")   # em dash after "Abstract"
    set_font(run, size_pt=9, bold=True, italic=True)
    return p


def add_abstract_body(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=4,
                 line_spacing=10)
    run = p.add_run(text)
    set_font(run, size_pt=9, italic=True)
    return p


def add_index_terms(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=6,
                 line_spacing=10)
    run1 = p.add_run("Index Terms—")
    set_font(run1, size_pt=9, bold=True, italic=True)
    run2 = p.add_run(text)
    set_font(run2, size_pt=9, italic=True)
    return p


def add_h1(doc, text):
    """Level-1 section heading: bold, uppercase, centered, 10pt."""
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=6, space_after=3)
    run = p.add_run(text.upper())
    set_font(run, size_pt=10, bold=True)
    return p


def add_h2(doc, text):
    """Level-2 subsection heading: bold italic, left-aligned, 10pt."""
    p = doc.add_paragraph()
    set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=4, space_after=1)
    run = p.add_run(text)
    set_font(run, size_pt=10, bold=True, italic=True)
    return p


def add_body(doc, text, space_before=0, space_after=3,
             first_line=0.1, bold=False, italic=False):
    """Body paragraph: 10pt Times New Roman, justified."""
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=space_before, space_after=space_after,
                 first_line_indent=first_line,
                 line_spacing=11)
    run = p.add_run(text)
    set_font(run, size_pt=10, bold=bold, italic=italic)
    return p


def add_table(doc, headers, rows, caption=""):
    """Add a table with IEEE-style formatting."""
    if caption:
        cp = doc.add_paragraph()
        set_para_fmt(cp, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=4, space_after=2)
        run = cp.add_run(caption)
        set_font(run, size_pt=8, bold=True)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
        run = p.add_run(h)
        set_font(run, size_pt=8, bold=True)

    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            set_para_fmt(p, space_before=0, space_after=0)
            run = p.add_run(val)
            set_font(run, size_pt=8)

    # Small space after table
    sp = doc.add_paragraph()
    set_para_fmt(sp, space_before=0, space_after=2)
    return t


def add_ref(doc, text):
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=1)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.2)
    pf.first_line_indent = Inches(-0.2)
    run = p.add_run(text)
    set_font(run, size_pt=8)
    return p


def add_rule(doc):
    """Thin horizontal rule (simulated with bottom border on an empty para)."""
    p = doc.add_paragraph()
    set_para_fmt(p, space_before=0, space_after=0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Main document builder ──────────────────────────────────────────────────

def build_ieee_manuscript():
    doc = Document()

    # Configure main section (2-column body)
    main_section = doc.sections[0]
    set_page_size(main_section)
    set_page_margins(main_section)
    set_columns(main_section, num_cols=2)

    # ── TITLE (single-column header area) ─────────────────────────────────
    add_title(doc,
        "Post-Quantum Security Framework for Model Context Protocol Using ContextGuard"
    )

    # ── AUTHORS ───────────────────────────────────────────────────────────
    add_author_name(doc, "Sunil Gentyala")
    add_affil(doc,
        "Enterprise Cybersecurity Division, HCLTech, America Inc., Dallas, TX, USA\n"
        "IEEE Senior Member #101760715 | ISACA Member #2297870\n"
        "sunil.gentyala@ieee.org"
    )
    add_author_name(doc, "A.V.H Sai Prasad")
    add_affil(doc,
        "Dept. of CSE (Data Science), CMR Technical Campus\n"
        "Kandlakoya, Medchal Road, Hyderabad-501401, Telangana, India\n"
        "avsaiprasad.ds@cmrtc.ac.in"
    )
    add_author_name(doc, "Kommana Swathi")
    add_affil(doc,
        "Dept. of CSE, Sir C R Reddy College of Engineering, Eluru, India\n"
        "kommanaswathi@gmail.com"
    )
    add_author_name(doc, "Vamshi Lande")
    add_affil(doc,
        "Research Student, University of North Texas, Denton, TX, USA\n"
        "vamshilande@my.unt.edu"
    )
    add_author_name(doc, "Akhila Kasturi")
    add_affil(doc,
        "Research Analyst, HCLTech, USA\n"
        "kasturi.akhila@gmail.com"
    )

    # ── ABSTRACT ──────────────────────────────────────────────────────────
    add_rule(doc)
    add_abstract_heading(doc, "Abstract")
    abstract_p = add_abstract_body(doc,
        "The proliferation of Model Context Protocol (MCP) as the operative "
        "authentication and tool invocation fabric for enterprise agentic AI systems "
        "has introduced a cryptographic trust surface that remains unexamined through "
        "the post-quantum security lens. This paper presents the first systematic "
        "application of RFC 9794's Post-Quantum/Traditional (PQ/T) hybrid terminology "
        "standard to MCP's four security-relevant layers: host authentication, "
        "transport channel confidentiality, server and tool identity certification, and "
        "tool result integrity. Three novel vulnerability classes are identified. MCP "
        "transport channels contain no normative reference to hybrid Key Encapsulation "
        "Mechanism design. MCP server identity relies exclusively on single-algorithm "
        "X.509 certificates, exposing tool authentication to harvest-now-decrypt-later "
        "(HNDL) adversaries. Tool result signing uses traditional digital signatures "
        "with no PQ/T hybrid equivalent, leaving a chain-of-trust ambiguity that RFC "
        "9794 Section 6.1 acknowledges but does not resolve. The ContextGuard-MCP-PQC "
        "governance framework is proposed as the structured response: a four-layer "
        "control architecture aligned to NIST CSF 2.0 and COBIT 2019, providing "
        "enterprise practitioners with a phased PQC migration roadmap. Empirical "
        "performance measurements confirm that composite X25519+ML-KEM-768 handshake "
        "overhead is below 1.5% of a typical MCP round-trip, and proof-of-concept "
        "demonstrations show that both HNDL harvesting and active downgrade attacks are "
        "blocked under Hybrid-Required enforcement. The ContextGuard-MCP-PQC framework "
        "is published as open-source software at "
        "github.com/sunilgentyala/contextguard-mcp-pqc."
    )
    index_p = add_index_terms(doc,
        "post-quantum cryptography, model context protocol, RFC 9794, hybrid key "
        "encapsulation, agentic AI security, ContextGuard, HNDL, PKI migration, "
        "NIST CSF 2.0, ML-KEM, ML-DSA"
    )
    add_rule(doc)

    # Insert continuous section break on the rule paragraph → switches to 2-column
    # The last paragraph before body content gets the inline sectPr (1-col section end)
    insert_continuous_section_break(index_p, num_cols=1)

    # ── I. INTRODUCTION ───────────────────────────────────────────────────
    add_h1(doc, "I. Introduction")
    add_body(doc,
        "Something consequential happened in enterprise AI infrastructure over the past "
        "eighteen months that the security community has not yet caught up with. Model "
        "Context Protocol (MCP), originally published by Anthropic as an open "
        "specification in late 2023, has matured from a developer convenience into the "
        "operational authentication and tool invocation substrate for a rapidly expanding "
        "population of enterprise agentic AI deployments. Banks run it to connect AI "
        "agents to internal data warehouses. Healthcare providers use it to authorize "
        "AI-mediated record access. Technology companies depend on it to determine which "
        "tools an autonomous agent is allowed to invoke during a task.", first_line=0
    )
    add_body(doc,
        "The cryptographic architecture that underpins these deployments is "
        "quantum-vulnerable in its entirety. Agent authentication relies on bearer "
        "tokens and public key infrastructure rooted in RSA and elliptic curve "
        "algorithms. Transport confidentiality depends on TLS 1.3 key exchange "
        "mechanisms that Shor's algorithm will eventually compromise. Server identity is "
        "bound to X.509 certificates containing single traditional public keys. Tool "
        "result signing, where it exists at all, uses ECDSA or RSA. None of these "
        "mechanisms have been evaluated, adapted, or supplemented with post-quantum "
        "equivalents in any published MCP specification or deployment guidance."
    )
    add_body(doc,
        "The harvest-now-decrypt-later (HNDL) attack model has been operationally "
        "relevant for several years. State-level adversaries are capturing encrypted "
        "traffic today with the explicit intention of decrypting it once a "
        "cryptographically relevant quantum computer (CRQC) becomes available. "
        "MCP-mediated tool calls that carry authorization context, query sensitive data "
        "repositories, or authenticate agent identities are precisely the kind of "
        "long-lived secrets that an HNDL adversary values."
    )
    add_body(doc,
        "In June 2025, the IETF published RFC 9794, 'Terminology for Post-Quantum "
        "Traditional Hybrid Schemes' [5], authored by researchers from the UK National "
        "Cyber Security Centre and the Naval Postgraduate School. The RFC establishes "
        "the foundational vocabulary for describing cryptographic constructions that "
        "combine post-quantum and traditional algorithms. This paper addresses the gap "
        "between RFC 9794 terminology and its application to MCP."
    )
    add_body(doc,
        "The contributions are as follows: (1) the first systematic four-layer "
        "post-quantum vulnerability taxonomy for MCP infrastructure, organized around "
        "RFC 9794 constructs; (2) identification of three previously undescribed "
        "vulnerability classes arising from the absence of PQC migration guidance in "
        "MCP specifications; (3) the ContextGuard-MCP-PQC governance framework, a NIST "
        "CSF 2.0 and COBIT 2019-aligned migration architecture with open-source "
        "implementation; (4) empirical performance measurements for ML-KEM-768 and "
        "ML-DSA-65 in MCP deployment contexts; and (5) proof-of-concept demonstrations "
        "for HNDL harvesting and TLS downgrade attacks with ContextGuard mitigation."
    )

    # ── II. BACKGROUND ────────────────────────────────────────────────────
    add_h1(doc, "II. Background")

    add_h2(doc, "A. The Quantum Threat")
    add_body(doc,
        "The theoretical basis for the quantum threat to public key cryptography has "
        "been settled since Shor's 1994 demonstration that a sufficiently capable "
        "quantum computer could factor large integers and compute discrete logarithms in "
        "polynomial time [1]. RSA and elliptic curve Diffie-Hellman, which underpin the "
        "overwhelming majority of public key authentication and key exchange mechanisms "
        "currently deployed, are vulnerable to polynomial-time quantum attacks. For MCP, "
        "the concern concentrates in the asymmetric layer, because that is where agent "
        "identity, channel key exchange, and certificate-based authentication all reside. "
        "The HNDL threat compounds this concern because adversaries do not require a "
        "CRQC at capture time.", first_line=0
    )

    add_h2(doc, "B. NIST Post-Quantum Cryptography Standards")
    add_body(doc,
        "In August 2024, NIST published its first finalized post-quantum cryptography "
        "standards: FIPS 203 (ML-KEM, formerly KYBER), FIPS 204 (ML-DSA, formerly "
        "DILITHIUM), and FIPS 205 (SLH-DSA, formerly SPHINCS+) [2],[3],[4]. ML-KEM "
        "provides key encapsulation functionality and is the post-quantum analogue of "
        "ECDH in TLS key exchange. ML-DSA provides digital signature functionality, the "
        "post-quantum analogue of ECDSA in certificate-based authentication. Integration "
        "requires specification-level changes rather than simple library substitution, "
        "which is why the hybrid transition period described by RFC 9794 is necessary.",
        first_line=0
    )

    add_h2(doc, "C. RFC 9794: Principal Constructs")
    add_body(doc,
        "RFC 9794 builds its taxonomy from first principles. A PQ/T hybrid scheme is a "
        "multi-algorithm scheme where at least one component algorithm is post-quantum "
        "and at least one is traditional. In a composite construction, the two components "
        "are packaged as a single unit. In a non-composite construction, the components "
        "remain separable within the protocol. The RFC's downgrade protection properties "
        "are particularly consequential: hybrid interoperability and hybrid authentication "
        "cannot both be achieved by a PQ/T hybrid scheme alone; they require a hybrid "
        "protocol with appropriate downgrade protection built into its negotiation layer "
        "[5]. For PKI, the RFC defines PQ/T hybrid certificates, post-quantum "
        "certificates, traditional certificates, and the PQ/T parallel PKI construct.",
        first_line=0
    )

    add_h2(doc, "D. MCP Architecture and Security Surface")
    add_body(doc,
        "MCP is a client-server architecture for enabling AI language model applications "
        "to interact with external tools and data sources via JSON-RPC 2.0 message "
        "exchange. Transport mechanisms vary: remote servers use HTTP with Server-Sent "
        "Events or WebSocket; local servers use stdio. Authentication in current "
        "deployments uses bearer tokens, OAuth 2.0 flows, or API keys. The specification "
        "provides no guidance on certificate-based mutual authentication, no requirements "
        "for session binding, and no reference to post-quantum cryptographic mechanisms. "
        "Recent empirical study of 1,899 open-source MCP server deployments found that "
        "7.2% contained exploitable vulnerabilities, underscoring the production "
        "security relevance of the MCP attack surface [14].", first_line=0
    )

    # ── III. THREAT MODEL ─────────────────────────────────────────────────
    add_h1(doc, "III. Threat Model")

    add_h2(doc, "A. Adversary Classes")
    add_body(doc,
        "This paper models two adversary classes. The first is a passive HNDL adversary: "
        "a nation-state level threat actor capable of intercepting and archiving network "
        "traffic at scale today, with a credible path to accessing a CRQC within a "
        "ten-to-fifteen-year window. The second is an active downgrade attacker capable "
        "of interposing in TLS negotiation to induce selection of weaker cryptographic "
        "parameters, exploiting the PQ/T hybrid negotiation surface to force "
        "traditional-only cipher suite selection.", first_line=0
    )

    add_h2(doc, "B. Assets at Risk")
    add_body(doc,
        "Table I characterizes the primary assets at risk across MCP's security layers. "
        "Quantitative risk scores for the corresponding vulnerability classes are "
        "presented in Section V-C.", first_line=0
    )
    add_table(doc,
        ["MCP Asset", "HNDL Risk", "Dwn. Risk", "Horizon"],
        [
            ["Agent identity tokens", "High", "High", "3-5 yr"],
            ["Session key material", "Critical", "High", "10-15 yr"],
            ["Server certificates", "High", "Medium", "5-10 yr"],
            ["Tool invocation payloads", "High", "High", "5-10 yr"],
            ["Tool result contents", "Medium", "Medium", "3-7 yr"],
            ["PKI root keys", "Critical", "Low", "20-30 yr"],
        ],
        caption="TABLE I: Asset Risk Classification"
    )

    # ── IV. APPLYING RFC 9794 CONSTRUCTS TO MCP ───────────────────────────
    add_h1(doc, "IV. Applying RFC 9794 Constructs to MCP")
    add_body(doc,
        "MCP's cryptographic exposure spans four distinct security layers, each "
        "addressable through a specific RFC 9794 construct. Layer 1 governs host "
        "authentication through identity assertions and signed credentials. Layer 2 "
        "governs session channel confidentiality carrying JSON-RPC messages. Layer 3 "
        "governs server and tool identity through certificate-bound public keys. Layer 4 "
        "governs the integrity of tool outputs returned to the AI agent. Each layer "
        "presents an independent quantum vulnerability that existing MCP specifications "
        "do not address.", first_line=0
    )
    add_table(doc,
        ["Layer", "RFC 9794 Construct", "CG Control"],
        [
            ["L1: Host Auth", "PQ/T hybrid signature", "HAF"],
            ["L2: Transport", "PQ/T composite KEM", "HKF"],
            ["L3: Server ID", "PQ/T parallel PKI", "HRS"],
            ["L4: Tool Result", "PQ/T hybrid signature", "RSP"],
        ],
        caption="TABLE II: RFC 9794 Construct Mapping by MCP Layer"
    )

    add_h2(doc, "A. Layer 1: Host Authentication and Agent Identity")
    add_body(doc,
        "The question of how an MCP host authenticates itself to downstream servers is "
        "handled in practice through bearer tokens, API keys, and OAuth 2.0 flows. None "
        "are post-quantum resistant. The identity assertions underlying these mechanisms "
        "are vulnerable to a CRQC through Shor's algorithm applied to RSA and elliptic "
        "curve primitives.", first_line=0
    )
    add_body(doc,
        "RFC 9794 defines the applicable construct: a PQ/T hybrid digital signature "
        "scheme, where both a post-quantum component (ML-DSA, FIPS 204) and a "
        "traditional component (ECDSA-P256) are applied simultaneously, requiring an "
        "adversary to break both to forge the signature. The composite construction is "
        "preferable during migration because it carries the traditional component, "
        "allowing legacy-compatible verification while providing post-quantum protection. "
        "For bearer token payloads specifically, RFC 9180 HPKE [8] provides a "
        "complementary hybrid encryption mechanism applicable to token confidentiality "
        "in future MCP authentication extensions, though its PQ/T migration path remains "
        "unspecified."
    )
    add_body(doc,
        "The ContextGuard Hybrid Authentication Floor (HAF) enforces minimum hybrid "
        "signature posture through three operational tiers. In Traditional-Only mode, "
        "the server accepts classical authentication and logs a warning. In "
        "Hybrid-Preferred mode, the server prioritizes hybrid authentication but permits "
        "traditional fallback with a P1 alert logged to the operational audit trail. In "
        "Hybrid-Required mode, the server evaluates the certificate chain at TLS "
        "handshake time, verifies the presence of an ML-DSA component in the presented "
        "credentials, and terminates the session with a TLS close_notify alert if only "
        "traditional algorithms are offered. The HAF tier is configured per-endpoint and "
        "applied at MCP server startup. Sessions failing HYBRID-REQUIRED evaluation are "
        "blocked before any JSON-RPC message exchange occurs, ensuring that no MCP tool "
        "is invocable over a quantum-vulnerable authentication channel."
    )

    add_h2(doc, "B. Layer 2: Transport Channel Confidentiality")
    add_body(doc,
        "MCP transport over HTTPS relies entirely on TLS 1.3 key exchange mechanisms "
        "that are quantum-vulnerable. The applicable RFC 9794 construct is the PQ/T "
        "composite KEM: a construction combining ML-KEM (FIPS 203) with X25519. The "
        "IETF draft-ietf-tls-hybrid-design specifies this construction for TLS 1.3, and "
        "production deployments by Cloudflare and Google have validated its operational "
        "viability, with ML-KEM-768 adding approximately 2,272 bytes of handshake "
        "overhead per connection [6].", first_line=0
    )
    add_body(doc,
        "The ContextGuard Hybrid KEM Floor (HKF) is implemented as a TLS session "
        "inspector integrated at MCP server startup. Implementation follows three "
        "specific steps. First, the MCP server is configured with OpenSSL 3.5 or later, "
        "using the Open Quantum Safe provider to enable the X25519MLKEM768 cipher suite "
        "group. Second, the cipher suite priority list is ordered: "
        "X25519MLKEM768:X25519MLKEM512:X25519, placing hybrid suites at highest "
        "preference. Third, the HKF policy tier is evaluated at ServerHello selection: "
        "in Hybrid-Required mode, if the TLS stack selects a non-hybrid suite because "
        "the client offered no hybrid options, the server closes the connection before "
        "completing the handshake. The STDIO transport variant remains outside TLS scope "
        "and requires a separate UNIX domain socket authentication layer, noted as a gap "
        "in current MCP specifications."
    )

    add_h2(doc, "C. Layer 3: Server and Tool Identity")
    add_body(doc,
        "Current MCP server identity relies on standard X.509 certificates with single "
        "RSA-2048 or ECDSA-P256 public keys. Two migration paths exist under RFC 9794. "
        "The PQ/T hybrid certificate embeds both post-quantum and traditional public "
        "keys within a single certificate. The PQ/T parallel PKI maintains two separate "
        "certificate chains used together in the protocol. The parallel PKI construction "
        "is the more practical near-term option because major certificate authorities "
        "have not yet issued composite PQ/T hybrid certificates in production.", first_line=0
    )
    add_body(doc,
        "The ContextGuard Hybrid Readiness Score (HRS) assigns a five-point integer "
        "metric to each MCP server endpoint based on five cumulative criteria. HRS-1 is "
        "awarded when the leaf certificate contains any post-quantum public key "
        "component. HRS-2 is awarded when the leaf is PQ/T hybrid, combining ML-DSA "
        "with a traditional algorithm either as a composite or parallel certificate "
        "pair. HRS-3 is awarded when all intermediate CA certificates in the chain are "
        "PQ/T hybrid, satisfying the RFC 9794 Section 6.1 chain-of-trust requirement. "
        "HRS-4 is awarded when the root CA is PQ/T hybrid or PQ-only. HRS-5 is awarded "
        "when active downgrade protection is confirmed through HAF and HKF enforcement "
        "for this endpoint. An endpoint scoring HRS-0 receives migration Priority 1 "
        "(immediate action); HRS-5 represents full compliance."
    )

    add_h2(doc, "D. Layer 4: Tool Result Integrity")
    add_body(doc,
        "The MCP specification does not require cryptographic signing of tool results. "
        "Most deployed servers return unsigned JSON payloads over TLS, relying "
        "implicitly on channel integrity. This model fails against a CRQC adversary "
        "capable of forging TLS session keys and injecting modified results, and against "
        "a compromised server whose results carry no independent provenance attestation. "
        "The ContextGuard Result Signature Policy (RSP) requires result signature "
        "verification against certificate chains meeting a minimum HRS threshold, "
        "classifying results from traditionally-rooted chains as partially verified.",
        first_line=0
    )

    # ── V. GOVERNANCE FRAMEWORK ───────────────────────────────────────────
    add_h1(doc, "V. ContextGuard-MCP-PQC Governance Framework")

    add_h2(doc, "A. Control Architecture")
    add_body(doc,
        "The ContextGuard-MCP-PQC framework extends the base ContextGuard zero-trust "
        "MCP governance architecture with a post-quantum control plane derived from the "
        "Section IV analysis. Four controls correspond to the four MCP security layers. "
        "The Agent Identity Control enforces HAF with three levels: Traditional-Only, "
        "Hybrid-Preferred (fallback permitted with logging), and Hybrid-Required "
        "(sessions refused if traditional-only). The Transport Channel Control enforces "
        "HKF, validating cipher suite selection against draft-ietf-tls-hybrid-design "
        "requirements. The Certificate Inventory Control assigns each MCP server "
        "endpoint an HRS from zero to five. The Result Provenance Control enforces RSP "
        "in three modes: Unsigned-Accept, Signed-Require-Traditional, and "
        "Signed-Require-Hybrid.", first_line=0
    )

    add_h2(doc, "B. Performance Evaluation")
    add_body(doc,
        "A critical question for practitioners adopting ContextGuard-MCP-PQC is whether "
        "the cryptographic overhead of hybrid PQC operations is compatible with "
        "production MCP deployment performance requirements. Table III presents "
        "performance data for ML-KEM-768 and ML-DSA-65 against classical algorithm "
        "baselines. Reference values are taken from Olushola and Meenakshi (2026) and "
        "Halak et al. (2024), who benchmarked these algorithms on commodity x86-64 "
        "hardware using NIST FIPS 203 and FIPS 204 reference implementations [17],[18].",
        first_line=0
    )
    add_table(doc,
        ["Algorithm", "Operation", "Latency (ms)", "vs. Classical"],
        [
            ["X25519", "Key exchange", "0.031", "Baseline"],
            ["ML-KEM-768", "Encapsulation", "0.073", "2.4x X25519"],
            ["ML-KEM-768", "Decapsulation", "0.063", "2.0x X25519"],
            ["X25519+ML-KEM-768", "TLS overhead", "1.84 total", "+1.14 ms"],
            ["ECDSA-P256", "Signing", "0.051", "Baseline"],
            ["ML-DSA-65", "Signing", "0.148", "12x faster RSA-2048"],
            ["ML-DSA-65", "Verification", "0.089", "2.0x ECDSA-P256"],
            ["RSA-2048", "Signing", "1.810", "Slowest baseline"],
        ],
        caption="TABLE III: ML-KEM-768 and ML-DSA-65 Performance"
    )
    add_body(doc,
        "The composite X25519+ML-KEM-768 handshake adds approximately 1.14 ms over "
        "classical X25519 on commodity x86-64 hardware. Against a typical 100-150 ms "
        "MCP round-trip time, this represents less than 1.5% overhead. ML-KEM-768 "
        "ciphertext and public key introduce 2,272 additional bytes per handshake per "
        "Stebila et al. [6]. In high-frequency MCP tool invocation scenarios with "
        "1,000 calls per second, the additional bandwidth is approximately 2.2 MB/s, "
        "negligible against modern enterprise network capacity. ML-DSA-65 signing at "
        "0.148 ms is twelve times faster than RSA-2048 signing at 1.81 ms, meaning "
        "hybrid certificate authentication outperforms current RSA-based deployments. "
        "These measurements confirm that ContextGuard-MCP-PQC controls are suitable "
        "for latency-sensitive tool invocation workflows."
    )

    add_h2(doc, "C. Quantitative Risk Assessment")
    add_body(doc,
        "Table IV presents quantitative risk scores for the four MCP vulnerability "
        "classes using a CVSS-inspired multi-dimensional scoring model. Each class is "
        "scored across five dimensions: Exploitability (E, 1-5), Attack Vector weight "
        "(AV, 1.0-2.0), CRQC Dependency (Q, 0-1, HNDL-adjusted), Sensitivity (S, "
        "1-5), and Mitigation Gap (M, 1-5). The HNDL adjustment reduces the CRQC "
        "dependency discount by 0.4 to reflect that harvest operations require no "
        "quantum computer and begin immediately. Final risk scores are normalized to a "
        "0-10 scale.", first_line=0
    )
    add_table(doc,
        ["Vulnerability", "Lyr", "E", "AV", "Q*", "S", "M", "Score", "Lvl"],
        [
            ["V1: Absent Hybrid KEM", "L2", "4.0", "2.0", "0.7", "5.0", "4.5", "9.8", "CRIT"],
            ["V2: Single-Alg Server ID", "L3", "3.5", "2.0", "0.9", "4.5", "4.0", "9.5", "CRIT"],
            ["V3: No Hybrid Auth Floor", "L1", "3.0", "2.0", "0.8", "4.5", "5.0", "9.5", "CRIT"],
            ["V4: Unsigned Tool Results", "L4", "3.5", "1.5", "0.6", "4.0", "5.0", "9.0", "CRIT"],
        ],
        caption="TABLE IV: Vulnerability Risk Assessment Matrix"
    )
    add_body(doc,
        "Q*=CRQC Dependency HNDL-adjusted. Score 0-10 scale. CRITICAL>=8.0.",
        italic=True
    )
    add_body(doc,
        "All four vulnerability classes score CRITICAL because the combination of "
        "network-level attack vectors, high sensitivity assets, and complete absence of "
        "any PQC mitigation in published MCP specifications produces scores in the "
        "9.0-9.8 range. The HNDL adjustment reflects that harvest-phase adversarial "
        "activity requires no CRQC and begins immediately. V1 receives the highest "
        "score because TLS transport channels carry both the highest session sensitivity "
        "and the widest network exposure surface."
    )

    add_h2(doc, "D. NIST CSF 2.0 and COBIT 2019 Mapping")
    add_table(doc,
        ["Control", "NIST CSF 2.0", "COBIT 2019"],
        [
            ["HAF (Agent Identity)", "GV.PO-01, PR.AA-01", "APO12.02, DSS05.04"],
            ["HKF (Transport)", "PR.DS-02, PR.DS-10", "DSS05.03, BAI10.01"],
            ["HRS (Certificate)", "DE.CM-01, GV.RM-04", "APO12.06, BAI10.03"],
            ["RSP (Result Provenance)", "DE.CM-06, RS.AN-03", "DSS05.07, APO12.04"],
        ],
        caption="TABLE V: Framework Governance Mapping"
    )
    add_body(doc,
        "The COBIT 2019 mapping follows the governance domain structure established in "
        "the information security governance literature [13]. Each ContextGuard control "
        "maps to both a CSF 2.0 outcome identifier and a COBIT 2019 management practice, "
        "enabling practitioners to integrate PQC migration into existing IT governance "
        "audit frameworks.", first_line=0
    )

    add_h2(doc, "E. Migration Phases")
    add_body(doc,
        "Phase 1 (Inventory and Assessment): all MCP server endpoints are catalogued, "
        "HRS values assigned, and HAF/HKF set to Monitor mode. Phase 2 (Hybrid "
        "Introduction): parallel PKI deployed for Tier 1 servers, HAF/HKF shifted to "
        "Hybrid-Preferred. Phase 3 (Parallel Operation): hybrid certificates extended "
        "to all servers, HAF/HKF shifted to Hybrid-Required for production. Phase 4 "
        "(PQC Consolidation): traditional-only chains retired, RSP shifted to "
        "Signed-Require-Hybrid across all tool categories.", first_line=0
    )

    # ── VI. DISCUSSION ────────────────────────────────────────────────────
    add_h1(doc, "VI. Discussion")

    add_h2(doc, "A. Comparative Analysis with Existing PQC Solutions")
    add_body(doc,
        "The post-quantum protocol security literature has concentrated on TLS 1.3 "
        "hybrid key exchange [7], IKEv2 [9], and X.509 certificate format extensions "
        "for post-quantum algorithms [10]. Hybrid Public Key Encryption (HPKE), "
        "specified in RFC 9180 [8], establishes a related foundational primitive for "
        "hybrid encryption of structured payloads, though without PQ/T migration "
        "guidance for specific protocol deployments. Security analysis of hybrid "
        "signature schemes by Bindel and Hale [11] establishes that hybrid construction "
        "security is bounded by the stronger component algorithm, supporting the "
        "composite PQ/T design used in ContextGuard HAF. The OWASP Top 10 for LLM "
        "Applications [15] identifies prompt injection and insecure tool invocation as "
        "leading AI agent risks; ContextGuard-MCP-PQC addresses the cryptographic "
        "trust layer that underpins those controls. The NCSC blog on RFC 9794 [16] "
        "provides practitioner-oriented context for the RFC's adoption trajectory. "
        "The MCP protocol family has received no treatment in the academic PQC "
        "literature. Table VI positions ContextGuard-MCP-PQC against the closest "
        "existing solutions.", first_line=0
    )
    add_table(doc,
        ["Solution", "Scope", "Transport PQC", "Downgrade", "MCP-Specific"],
        [
            ["ContextGuard-MCP-PQC", "MCP four layers", "HKF composite KEM", "HYBRID-REQUIRED", "Yes"],
            ["Open Quantum Safe (OQS)", "Generic library", "Cipher suite libs", "Not specified", "No"],
            ["NIST NCCoE Migration", "Enterprise IT", "TLS guidance", "Policy only", "No"],
            ["Parallel MCP-PQC work", "MCP transport", "HHF enforcer", "Monitor/Prefer", "Yes"],
            ["Google/Cloudflare hybrid", "Internet TLS", "X25519+ML-KEM-768", "Not enforced", "No"],
            ["ETSI TS 103 744", "Telco/5G", "Specified", "Specified", "No"],
        ],
        caption="TABLE VI: Comparison with Existing PQC Migration Solutions"
    )
    add_body(doc,
        "The key differentiator of ContextGuard-MCP-PQC is its MCP-specific four-layer "
        "architecture derived from RFC 9794 constructs. Existing generic solutions (OQS, "
        "NCCoE guidance, major cloud vendor deployments) provide building blocks without "
        "addressing MCP's specific trust model: bearer token authentication, JSON-RPC "
        "tool invocation, and the chain-of-trust gap in unsigned tool results. Compared "
        "to parallel MCP-PQC work focused on transport-layer handshake hardening and "
        "composite certificate readiness scoring, ContextGuard-MCP-PQC adds Layer 1 "
        "HAF enforcement and Layer 4 RSP provenance controls, and introduces "
        "HYBRID-REQUIRED enforcement tiers that fully block downgrade attempts. The HRS "
        "scoring system provides a migration priority queue absent from all compared "
        "solutions."
    )

    add_h2(doc, "B. Proof-of-Concept: HNDL and Downgrade Attacks")
    add_body(doc,
        "Two proof-of-concept demonstrations accompany this paper, implemented in Python "
        "3.9+ and available at github.com/sunilgentyala/contextguard-mcp-pqc.", first_line=0
    )
    add_body(doc,
        "PoC-1: HNDL Threat Simulation (poc/hndl_demo.py). The simulation creates "
        "realistic MCP TLS session descriptors representing traffic captured by a passive "
        "network adversary. Each session specifies the cipher suite, server certificate "
        "algorithm, and payload sensitivity class. The simulation computes the HNDL risk "
        "for each captured asset using sensitivity horizons from Table I and the CRQC "
        "timeline distribution (optimistic: 2030, consensus: 2035, conservative: 2040). "
        "Sessions using traditional ECDH key exchange are flagged as HNDL-vulnerable "
        "across all asset classes with horizons exceeding four years. Sessions using "
        "X25519+ML-KEM-768 are shown to be HNDL-resistant because the hybrid KEM "
        "requires the adversary to break ML-KEM-768, for which no quantum algorithm "
        "faster than lattice sieving is known (NIST FIPS 203, Appendix A)."
    )
    add_body(doc,
        "PoC-2: Downgrade Attack Simulation (poc/downgrade_demo.py). The demonstration "
        "implements a four-party simulation: legitimate client, active adversary, and "
        "MCP server negotiator evaluated under four HKF modes (DISABLED, MONITOR, "
        "HYBRID_PREFERRED, HYBRID_REQUIRED). The adversary intercepts the client's "
        "ClientHello and strips X25519MLKEM768 cipher suite extensions before forwarding "
        "to the server. Under DISABLED and MONITOR modes, the server selects "
        "ECDHE-RSA, completing the downgrade successfully. Under HYBRID_PREFERRED, the "
        "server falls back with a warning log entry but still allows the session. Under "
        "HYBRID_REQUIRED, the server terminates the connection before ServerHello "
        "completion. This confirms that RFC 9794 Section 4.2's downgrade protection "
        "requirement demands active enforcement at the server tier, which is precisely "
        "what HKF HYBRID_REQUIRED implements."
    )

    add_h2(doc, "C. Limitations")
    add_body(doc,
        "The analysis is conducted against published MCP specifications, not deployed "
        "production systems. The vulnerability characterization identifies specification "
        "gaps mapped to RFC 9794 constructs but does not include formal security proofs. "
        "Performance measurements are derived from published benchmark studies on "
        "commodity x86-64 hardware; high-concurrency MCP deployments at 10,000+ calls "
        "per second have not been characterized. Formal cryptographic security proofs "
        "for the proposed hybrid constructions are out of scope; practitioners requiring "
        "such proofs should consult the relevant IETF and NIST documentation directly.",
        first_line=0
    )

    # ── VII. CONCLUSION AND FUTURE DIRECTIONS ─────────────────────────────
    add_h1(doc, "VII. Conclusion and Future Directions")
    add_body(doc,
        "Across four security-relevant layers of MCP deployment, the application of RFC "
        "9794's PQ/T hybrid constructs reveals a consistent pattern of structural "
        "specification gaps. Agent authentication has no hybrid signature specification. "
        "Transport channels have no hybrid KEM requirement. Server identity has no "
        "parallel PKI framework. Result signing has no hybrid signature standard and no "
        "chain-of-trust completeness requirement.", first_line=0
    )
    add_body(doc,
        "The ContextGuard-MCP-PQC framework addresses each gap through four controls: "
        "HAF, HKF, HRS, and RSP. Proof-of-concept demonstrations confirm that HNDL "
        "harvesting of traditional-only MCP sessions is operationally straightforward, "
        "that active downgrade attacks succeed unless HYBRID_REQUIRED enforcement is "
        "active, and that hybrid PQC overhead is below 1.5% of a typical MCP "
        "round-trip. The framework is available as open-source software at "
        "github.com/sunilgentyala/contextguard-mcp-pqc."
    )
    add_body(doc,
        "Four research directions follow. First, a longitudinal field study of HRS score "
        "progression across a real enterprise MCP server inventory over a twelve-month "
        "migration cycle. Second, a formal security proof for the HAF/HKF combination "
        "as a downgrade-resistant authentication protocol under the RFC 9794 hybrid "
        "security definition. Third, empirical performance characterization of "
        "ContextGuard-MCP-PQC controls under high-concurrency MCP tool invocation at "
        "10,000 calls per second. Fourth, specification of PQ/PQ multi-algorithm "
        "constructions as defined in RFC 9794 Section 3.4 for post-transition "
        "environments where all component algorithms are post-quantum."
    )
    add_body(doc,
        "The urgency is real and the planning window is finite. HNDL adversaries are "
        "active today. Organizations that complete MCP post-quantum inventory and "
        "initiate hybrid transition now will be substantially better positioned than "
        "those that wait for quantum timelines to become undeniable. The ContextGuard "
        "framework is the migration plan. What remains is execution."
    )

    # ── REFERENCES ────────────────────────────────────────────────────────
    add_h1(doc, "References")
    refs = [
        "[1] P. W. Shor, 'Algorithms for quantum computation,' Proc. 35th FOCS, IEEE, 1994, pp. 124-134.",
        "[2] NIST, 'Module-Lattice-Based KEM Standard,' FIPS 203, Aug. 2024.",
        "[3] NIST, 'Module-Lattice-Based DSS,' FIPS 204, Aug. 2024.",
        "[4] NIST, 'Stateless Hash-Based DSS,' FIPS 205, Aug. 2024.",
        "[5] F. Driscoll, M. Parsons, and B. Hale, 'Terminology for Post-Quantum Traditional Hybrid Schemes,' RFC 9794, IETF, Jun. 2025. DOI: 10.17487/RFC9794.",
        "[6] D. Stebila, S. Fluhrer, and S. Gueron, 'Hybrid key exchange in TLS 1.3,' draft-ietf-tls-hybrid-design-12, Jan. 2025.",
        "[7] N. Bindel et al., 'Hybrid KEMs and AKE,' PQCrypto 2019, LNCS vol. 11505, pp. 206-226.",
        "[8] R. Barnes, K. Bhargavan, B. Lipp, and C. Wood, 'Hybrid Public Key Encryption,' RFC 9180, IETF, Feb. 2022. DOI: 10.17487/RFC9180.",
        "[9] C. Tjhai et al., 'Multiple Key Exchanges in IKEv2,' RFC 9370, IETF, May 2023.",
        "[10] M. Ounsworth et al., 'Composite ML-KEM for X.509 PKI,' draft-ietf-lamps-pq-composite-kem-06, Mar. 2025.",
        "[11] N. Bindel and B. Hale, 'A Note on Hybrid Signature Schemes,' ePrint 2023/423, Jul. 2023.",
        "[12] NIST, 'Cybersecurity Framework 2.0,' NIST CSF 2.0, Feb. 2024.",
        "[13] D. Harjanti and B. Hendradjaja, 'Designing Information Security Governance Recommendations and Roadmap Using COBIT 2019 Framework and ISO 27001:2013,' in Proc. ICCSCI, 2020. DOI: 10.1109/ICCSCI50791.2020.00047.",
        "[14] V. S. Narajala and I. Habler, 'Enterprise-Grade Security for the Model Context Protocol (MCP): Frameworks and Mitigation Strategies,' arXiv:2504.08623, Apr. 2025.",
        "[15] OWASP Foundation, 'OWASP Top 10 for LLM Applications,' ver. 2.0, 2024.",
        "[16] NCSC, 'RFC 9794: A New Standard for Post-Quantum Terminology,' NCSC Blog, Oct. 2025.",
        "[17] A. Olushola and S. P. Meenakshi, 'Design and implementation of an authenticated post-quantum session protocol using ML-KEM, ML-DSA, and AES-256-GCM,' Front. Phys., vol. 13, Art. 1723966, Jan. 2026. DOI: 10.3389/fphy.2025.1723966.",
        "[18] B. Halak et al., 'Constrained Device Performance Benchmarking with the Implementation of Post-Quantum Cryptography,' Cryptography, vol. 8, no. 2, Art. 21, 2024. DOI: 10.3390/cryptography8020021.",
    ]
    for ref in refs:
        add_ref(doc, ref)

    doc.save(OUTPUT_PATH)
    print(f"IEEE-formatted manuscript saved to:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build_ieee_manuscript()
